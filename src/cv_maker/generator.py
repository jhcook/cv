
# Copyright 2026 Justin Cook
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""
Handles the generation of the final MS Word (DOCX) CV.
"""

import logging
from docx import Document
from docx.shared import Pt
from cv_maker.models import CVData

logger = logging.getLogger(__name__)

class CVGenerator:
    """
    Generates a styled DOCX resume from structured CVData.
    """
    def __init__(self, template_path: str = None, suggestions: list = None):
        self.template_used = False
        self.template_has_header = False
        self.styles = {
            'title': 'Title',
            'h1': 'Heading 1', 
            'h2': 'Heading 2',
            'body': 'Normal',
            'bullet': 'List Bullet'
        }
        
        # Buffers for assembly strategy
        self.injections = {}         # {target_element: [new_elements]}
        self.header_elements = []    # Elements to prepend
        self.section_map = {}        # {section_key: target_element}
        
        if template_path:
            self.template_path = template_path
            try:
                logger.info(f"Loading template: {template_path}")
                self.document = Document(template_path)
                self.template_used = True
                
                # Detect styles before clearing
                self._detect_template_styles()
                
                # Map sections for smart injection
                self._map_template_sections()
                
                # Check if the template header already has name/contact
                self._detect_header_content()
                
                self._clear_body_content()
                
                # Handle suggestions
                if suggestions and 'header' in suggestions:
                    self._clear_headers_footers()
                
                if suggestions and 'font' in suggestions:
                    self._setup_styles()
                    
            except Exception as e:
                logger.error(f"Error loading template: {e}. Falling back to default.")
                self.document = Document()
                self._setup_styles()
        else:
            self.document = Document()
            self._setup_styles()

    def _detect_template_styles(self):
        """
        Heuristically identifies styles based on visual properties of the template text.
        Also detects Title and Subtitle/Contact styles from the top of the document.
        """
        try:
            # Stats: style_name -> {count, score, is_bullet}
            style_stats = {}
            paragraphs = self.document.paragraphs
            
            # 1. Detect Title & Subtitle (First 2 paragraphs usually)
            if len(paragraphs) > 0:
                self.styles['title'] = paragraphs[0].style.name
                logger.debug(f"    > Detected Title Style: '{self.styles['title']}'")
            
            if len(paragraphs) > 1:
                # Often the contact info or role
                self.styles['subtitle'] = paragraphs[1].style.name
                logger.debug(f"    > Detected Subtitle/Contact Style: '{self.styles['subtitle']}'")
            else:
                 self.styles['subtitle'] = 'Normal'

            for p in paragraphs:
                text = p.text.strip()
                if not text: continue
                # ... rest of loop ...
                
                name = p.style.name
                if name not in style_stats:
                    style_stats[name] = {'count': 0, 'header_score': 0, 'bullet_score': 0}
                
                stats = style_stats[name]
                stats['count'] += 1
                
                # Heuristics for Header
                # 1. Short text (likely a title)
                if len(text) < 50:
                    stats['header_score'] += 1
                
                # 2. Uppercase (highly likely header)
                if text.isupper() and len(text) > 4:
                    stats['header_score'] += 3
                
                # 3. Bold (often header)
                # Note: This checks style definition or direct formatting if accessible
                if p.runs and p.runs[0].bold:
                    stats['header_score'] += 2
                    
                # 4. Keyword Boost (still useful)
                lower_text = text.lower()
                if any(x in lower_text for x in ['experience', 'education', 'skills', 'summary', 'projects']):
                    stats['header_score'] += 5

                # Heuristics for Bullet
                if p.style.name.lower().find('list') != -1 or p.style.name.lower().find('bullet') != -1:
                    stats['bullet_score'] += 10
                if text.startswith('•') or text.startswith('-') or text.startswith('➢'):
                    stats['bullet_score'] += 5

            # Analyze Results
            # Find best H1 (highest header_score but excluding Title-like if possible)
            # Find best Bullet (highest bullet_score)
            
            best_h1 = None
            max_h_score = 0
            
            best_bullet = None
            max_b_score = 0
            
            for name, stats in style_stats.items():
                # print(f"DEBUG STYLE '{name}': H={stats['header_score']}, B={stats['bullet_score']}")
                
                if stats['header_score'] > max_h_score:
                    # Avoid picking 'Title' style as H1 if logical
                    if 'title' not in name.lower(): 
                        max_h_score = stats['header_score']
                        best_h1 = name
                
                if stats['bullet_score'] > max_b_score:
                     max_b_score = stats['bullet_score']
                     best_bullet = name
            
            if best_h1:
                self.styles['h1'] = best_h1
                logger.info(f"    > Heuristic Header Detection: Using '{best_h1}' (Score: {max_h_score})")
            
            if best_bullet:
                self.styles['bullet'] = best_bullet
                logger.info(f"    > Heuristic Bullet Detection: Using '{best_bullet}' (Score: {max_b_score})")

        except Exception as e:
            logger.warning(f"Warning detecting styles: {e}")

    def _map_template_sections(self):
        """
        Scans the document for existing Section Headers (e.g. "Experience", "Education").
        Stores their locations so we can inject content there instead of appending.
        
        Returns:
            dict: { 'experience': {'type': 'p'/'tbl', 'element': obj}, ... }
        """
        self.section_map = {}
        
        # Keywords to look for
        keywords = {
            'experience': ['experience', 'employment', 'work history'],
            'education': ['education', 'academic', 'qualifications'],
            'projects': ['projects', 'technical', 'open source'],
            'summary': ['summary', 'profile', 'about me'],
            'skills': ['competencies', 'skills', 'technologies']
        }
        
        # 1. Scan Paragraphs
        for p in self.document.paragraphs:
            text = p.text.strip().lower()
            if not text: continue
            
            for section, keys in keywords.items():
                if any(k in text for k in keys) and len(text) < 50:
                    # Found a potential header
                    if section not in self.section_map:
                         self.section_map[section] = {'type': 'paragraph', 'object': p}
                         logger.debug(f"    > Mapped '{section}' to Paragraph: '{p.text[:20]}...'")

        # 2. Scan Tables (look in first row/cell usually)
        for tbl in self.document.tables:
            # Check first row, first cell? or iterate all cells?
            # Usually header is a distinct row or cell.
            # Let's check the whole table's first few rows.
            for row in tbl.rows[:3]: 
                for cell in row.cells:
                    text = cell.text.strip().lower()
                    if not text: continue
                    
                    for section, keys in keywords.items():
                        if any(k in text for k in keys) and len(text) < 50:
                            if section not in self.section_map:
                                # We map to the TABLE, but maybe we need the specific cell?
                                # Ideally we want to write to the *next* row or the *same* cell?
                                # Let's store the Table and the Cell.
                                self.section_map[section] = {'type': 'table', 'object': tbl, 'cell': cell}
                                logger.debug(f"    > Mapped '{section}' to Table Cell: '{cell.text[:20]}...'")

    def _detect_header_content(self):
        """
        Inspects the template's first-page header for existing name/contact content.
        If found, sets self.template_has_header = True so generate() skips
        duplicating name/title/contact in the body.
        """
        self.template_has_header = False
        self.header_texts = []
        try:
            for section in self.document.sections:
                header = section.header
                if header and not header.is_linked_to_previous:
                    for p in header.paragraphs:
                        text = p.text.strip()
                        if text:
                            self.header_texts.append(text)
                # Also check first_page_header (different first page)
                first_header = section.first_page_header
                if first_header:
                    for p in first_header.paragraphs:
                        text = p.text.strip()
                        if text:
                            self.header_texts.append(text)

            if self.header_texts:
                self.template_has_header = True
                logger.info(f"    > Template header detected with {len(self.header_texts)} text element(s): "
                            f"{', '.join(t[:30] for t in self.header_texts)}")
            else:
                logger.debug("    > No text content in template header.")
        except Exception as e:
            logger.warning(f"Warning detecting header content: {e}")

    def _clear_body_content(self):
        """
        Removes content while respecting the section_map.
        Strategy:
        1. If an element is part of the section_map, KEEP IT (and its parents).
        2. If it's a Table, we generally keep it if mapped.
        3. Clear *other* paragraphs to remove placeholder text.
        """
        # Collect mapped objects for easy checked
        mapped_objs = []
        for v in self.section_map.values():
            mapped_objs.append(v['object'])
            
        try:
            body = self.document.element.body
            for element in list(body):
                # Always keep SectPr
                if element.tag.endswith('sectPr'): 
                    print("DEBUG: Preserving sectPr")
                    continue
                
                # Check for Graphics (preserve layout)
                if 'w:drawing' in element.xml or 'w:pict' in element.xml:
                     print(f"DEBUG: Preserving Graphics in {element.tag}")
                     continue

                # Check if this element corresponds to a mapped object
                # python-docx objects wrap elements. 
                # We need to check if element IS the wrapped element.
                is_mapped = False
                for obj in mapped_objs:
                    if hasattr(obj, '_element') and obj._element == element:
                        is_mapped = True
                        break
                
                if is_mapped:
                    print(f"DEBUG: Preserving mapped section element: {element.tag}")
                    continue
                
                # If it's a Table, and we decided NOT to map it, do we remove it?
                # If the user has a complex layout, they might have specific tables for layout 
                # that DON'T contain headers (e.g. wrapper tables).
                # Removing UNMAPPED tables is risky if we want to "preserve layout".
                # BUT if we keep them, they contain placeholder text.
                
                # safe approach: If we found ANY mapped sections, assumes "strict injection" 
                # and remove everything else?
                # OR: Clear the *text* inside unmapped tables/paragraphs?
                
                # "It's placing all the content with their own headers below."
                # This implied we were NOT clearing enough (or clearing the wrong things).
                
                # If we mapped "Experience", we keep that Table.
                # If there's another table with "Lorem Ipsum", we should probably clear it.
                
                if element.tag.endswith('p') or element.tag.endswith('tbl'):
                    try:
                        # Extract all text from w:t elements
                        texts = element.xpath('.//w:t/text()')
                        full_text = "".join(texts)
                    except:
                        full_text = "?"
                    
                    print(f"DEBUG: Removing {element.tag} - {full_text[:30]}")
                    if element.getparent() == body:
                        body.remove(element)
                        if element in body:
                            print(f"DEBUG: FAILED TO REMOVE {element.tag}")
                        else:
                            print(f"DEBUG: Successfully removed {element.tag}")
                    else:
                         print(f"DEBUG: Parent mismatch. Parent: {element.getparent()}")
                        
        except Exception as e:
            logger.warning(f"Warning cleaning template body: {e}")
        
        # DEBUG SAVE
        self.document.save("DEBUG_OUTPUT.docx")
        print("DEBUG: Saved intermediate DEBUG_OUTPUT.docx")

    def _clear_headers_footers(self):
        """Clears content from headers and footers."""
        try:
            for section in self.document.sections:
                if section.header:
                    for element in list(section.header._element):
                         if element.tag.endswith('p') or element.tag.endswith('tbl'):
                             section.header._element.remove(element)
                if section.footer:
                     for element in list(section.footer._element):
                         if element.tag.endswith('p') or element.tag.endswith('tbl'):
                             section.footer._element.remove(element)
        except Exception as e:
            logger.warning(f"Warning cleaning headers/footers: {e}")

    def _inject_content_after(self, target_obj, content_func):
        """
        Executes a function that generates content, but ensures that content 
        is inserted immediately after the target_obj (Paragraph or Table).
        
        Args:
            target_obj: The python-docx object (Paragraph or Table) to insert after.
            content_func: A callback(doc) that adds content. 
                          We will capture the added content and move it? 
                          Or we manually construct elements?
        
        Complexity: python-docx `add_paragraph` always appends to end of body.
        Moving elements in XML is safer.
        
        Strategy:
        1. Record the initial length of body elements.
        2. Run content_func(self.document) (which appends to end).
        3. Identify the NEW elements.
        4. Move them to be siblings of target_obj.
        """
        # 1. Snapshot end of body
        body = self.document.element.body
        initial_count = len(body)
        
        # 2. Generate content (appended to end)
        content_func() # This uses self.document.add_paragraph etc.
        
        # 3. Find new elements
        # Note: We must fetch body again because it was modified
        body = self.document.element.body
        new_elements = list(body)[initial_count:]
        
    def _inject_content_after(self, target_obj, content_func):
        """
        Executes content_func() to generate new elements (appended to body).
        Then moves those new elements to self.injections buffer for later assembly.
        """
        # 1. Snapshot end of body
        try:
            body = self.document.element.body
            initial_count = len(body)
            
            # 2. Generate content (appended to end)
            content_func()
            
            # 3. Find new elements
            body = self.document.element.body # Re-fetch
            new_elements = list(body)[initial_count:]
            
            # 4. Store in buffer, remove from body
            if hasattr(target_obj, '_element'):
                target_element = target_obj._element
                
                print(f"DEBUG: Buffering {len(new_elements)} elements for {target_element.tag}")
                
                # Remove from body immediately so they don't interfere with next generation
                for elem in new_elements:
                    if elem.getparent() is not None:
                        elem.getparent().remove(elem)
                
                # Store references
                if target_element not in self.injections:
                    self.injections[target_element] = []
                self.injections[target_element].extend(new_elements)
                
        except Exception as e:
            logger.error(f"Error injecting content: {e}")

    def _prepend_content(self, content_func):
        """
        Executes content_func(), buffers new elements in self.header_elements.
        """
        # 1. Snapshot end of body
        body = self.document.element.body
        initial_count = len(body)
        
        # 2. Generate content (appended to end)
        content_func()
        
        # 3. Buffer and remove
        body = self.document.element.body
        new_elements = list(body)[initial_count:]
        
        print(f"DEBUG: Buffering {len(new_elements)} header elements")
        
        for elem in new_elements:
            if elem.getparent() is not None:
                elem.getparent().remove(elem)
        
        self.header_elements.extend(new_elements)

    def _assemble_document(self):
        """
        Reconstructs the document body from:
        - self.header_elements (First)
        - Existing body elements (Preserved templates)
        - self.injections (Interleaved after targets)
        """
        print("DEBUG: Assembling document...")
        body = self.document.element.body
        
        # 1. Start with header
        final_elements = []
        final_elements.extend(self.header_elements)
        
        # 2. Iterate through current body elements
        # Note: These are the preserved ones (since we removed generated ones)
        current_body_elements = list(body)
        
        for elem in current_body_elements:
            final_elements.append(elem)
            
            # Check if this element is a target for injection
            if elem in self.injections:
                print(f"DEBUG: Injecting {len(self.injections[elem])} buffered elements after {elem.tag}")
                final_elements.extend(self.injections[elem])
        
        # 3. Replace body content
        # Clear body (careful not to delete the elements we just collected!)
        body[:] = final_elements 
        
        print(f"DEBUG: Assembly complete. Total elements: {len(body)}")

    def _setup_styles(self):
        try:
            style = self.document.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
        except:
            pass

    def generate(self, data: CVData, output_filename: str):
        """
        Main entry point to generate the document.
        
        Args:
            data (CVData): The structured CV data.
            output_filename (str): The path to save the generated DOCX.
        """
        # --- HELPER FUNCTIONS FOR CONTENT GENERATION ---
        
        def add_header_content():
             # Title (Name)
            p = self.document.add_paragraph(data.name)
            p.style = self.styles.get('title', 'Title')
            
            # Sub-header lines (Role, Contact)
            subtitle_style = self.styles.get('subtitle', 'Normal')
            
            contact = self.document.add_paragraph(data.title)
            contact.style = subtitle_style
            if contact.runs:
                contact.runs[0].bold = True
            else:
                contact.add_run(data.title).bold = True
            
            details = self.document.add_paragraph(data.contact_info)
            details.style = subtitle_style

            self.document.add_paragraph()  # Spacer

        def add_summary():
             self.document.add_paragraph(data.executive_summary)

        def add_competencies():
            for category, skills in data.competencies:
                p = self.document.add_paragraph(style=self.styles['bullet'])
                p.add_run(category).bold = True
                p.add_run(f" {skills}")

        def add_experience():
            for job in data.experience:
                self.document.add_paragraph() # Top spacer for job
                
                # Company Line
                p = self.document.add_paragraph()
                p.add_run(job.company.upper()).bold = True
                p.add_run(f" | {job.location} | ")
                p.add_run(job.dates).italic = True
                p.paragraph_format.keep_with_next = True
                
                # Title Line
                p = self.document.add_paragraph()
                p.add_run(job.title).bold = True
                p.paragraph_format.keep_with_next = True
                
                if job.summary_italic:
                    p = self.document.add_paragraph(job.summary_italic)
                    p.italic = True
                    p.paragraph_format.widow_control = True

                # Bullets
                for title, desc in job.bullets:
                    p = self.document.add_paragraph(style=self.styles['bullet'])
                    p.add_run(title).bold = True
                    p.add_run(f" {desc}")
                    p.paragraph_format.widow_control = True

        def add_earlier_experience():
             for job in data.earlier_experience:
                self.document.add_paragraph() # Spacer
                
                # Title, Company Line (No dates)
                p = self.document.add_paragraph()
                p.add_run(f"{job.title}, {job.company}").bold = True
                p.paragraph_format.keep_with_next = True
                
                # Summary
                p = self.document.add_paragraph(job.summary)
                p.paragraph_format.widow_control = True

        def add_projects():
             p = self.document.add_paragraph('Visible at: github.com/username')
             p.italic = True 
             p.paragraph_format.keep_with_next = True

             for title, desc in data.projects:
                p = self.document.add_paragraph(style=self.styles['bullet'])
                p.add_run(title).bold = True
                p.add_run(f" {desc}")
                p.paragraph_format.widow_control = True

        def add_education():
            for edu in data.education:
                p = self.document.add_paragraph(edu)
                p.paragraph_format.widow_control = True
            
            if data.certifications:
                certs = self.document.add_paragraph()
                certs.add_run('Certifications: ').bold = True
                certs.add_run(data.certifications)
                certs.paragraph_format.widow_control = True

        # --- HEADER EXECUTION ---
        if self.template_has_header:
            logger.info("    > Skipping body header (template header has name/contact)")
        else:
            # Prepend header content to ensure it is at the TOP of the body
            # (before any preserved template sections)
            self._prepend_content(add_header_content)

        # --- EXECUTIVE SUMMARY ---
        if 'summary' in self.section_map:
             self._inject_content_after(self.section_map['summary']['object'], add_summary)
        else:
             p = self.document.add_paragraph('EXECUTIVE SUMMARY', style=self.styles['h1'])
             add_summary()

        # --- CORE COMPETENCIES ---
        if data.competencies:
            if 'skills' in self.section_map:
                self._inject_content_after(self.section_map['skills']['object'], add_competencies)
            else:
                p = self.document.add_paragraph('CORE COMPETENCIES', style=self.styles['h1'])
                add_competencies()
 
        # --- PROFESSIONAL EXPERIENCE ---
        if 'experience' in self.section_map:
            self._inject_content_after(self.section_map['experience']['object'], add_experience)
        else:
            p = self.document.add_paragraph('PROFESSIONAL EXPERIENCE', style=self.styles['h1'])
            p.paragraph_format.keep_with_next = True
            add_experience()

        if 'experience' not in self.section_map:
             self.document.add_paragraph() # Spacer only if we appended

        # --- EARLIER CAREER EXPERIENCE ---
        if data.earlier_experience:
            # We treat this as part of experience usually, but if there is a separate section?
            # Unlikely. We just append it if not mapped? 
            # Or if experience IS mapped, do we inject this too?
            # It's safer to just append it blindly for now unless we search for "Earlier".
            # If "Experience" was mapped, we injected into it. 
            # We should probably inject this AFTER the experience content?
            # This is tricky because we just lost reference to the injected content location.
            # But the 'add_experience' function injects everything.
            # Let's just append it for now as a fallback, or handle it nicely later.
            p = self.document.add_paragraph('EARLIER CAREER EXPERIENCE', style=self.styles['h1'])
            p.paragraph_format.keep_with_next = True
            add_earlier_experience()

        if 'experience' not in self.section_map:
             self.document.add_paragraph() # Spacer

        # --- PROJECTS ---
        if data.projects:
            if 'projects' in self.section_map:
                self._inject_content_after(self.section_map['projects']['object'], add_projects)
            else:
                p = self.document.add_paragraph('TECHNICAL PROJECTS & OPEN SOURCE', style=self.styles['h1'])
                p.paragraph_format.keep_with_next = True
                add_projects()

        # --- EDUCATION ---
        if 'education' in self.section_map:
            self._inject_content_after(self.section_map['education']['object'], add_education)
        else:
            p = self.document.add_paragraph('EDUCATION & CERTIFICATIONS', style=self.styles['h1'])
            p.paragraph_format.keep_with_next = True
            add_education()
 
        # --- ASSEMBLE ---
        self._assemble_document()
        
        # Save
        self.document.save(output_filename)
        logger.info(f"CV generated successfully: {output_filename}")

    def generate_cover_letter(self, data: CVData, letter_body: str, output_filename: str):
        """
        Generates a styled Cover Letter in a separate DOCX.
        Reuses detected styles for the header to ensure brand consistency.
        """
        
        # Create new document for CL (or we could use a fresh template instance if we wanted full properties)
        # To reuse styles properly, we should re-load the template if it exists.
        if self.template_used and hasattr(self, 'template_path'):
             cl_doc = Document(self.template_path)
        else:
             cl_doc = Document()
        
        # Clear everything in the new doc instance
        try:
            body = cl_doc.element.body
            for element in list(body):
                if element.tag.endswith('sectPr'): continue
                if element.tag.endswith('p') or element.tag.endswith('tbl'):
                    body.remove(element)
        except: pass

        # Clear CV-branded header from cover letter
        try:
            for section in cl_doc.sections:
                if section.header:
                    for element in list(section.header._element):
                        if element.tag.endswith('p') or element.tag.endswith('tbl'):
                            section.header._element.remove(element)
                # Also clear first-page header if present
                first_header = section.first_page_header
                if first_header:
                    for element in list(first_header._element):
                        if element.tag.endswith('p') or element.tag.endswith('tbl'):
                            first_header._element.remove(element)
            logger.debug("    > Cleared CV template header from cover letter")
        except Exception as e:
            logger.warning(f"Warning clearing cover letter header: {e}")

        # --- HEADER (Same as CV) ---
        # Title (Name)
        p = cl_doc.add_paragraph(data.name)
        p.style = self.styles.get('title', 'Title')
        
        # Contact Info
        subtitle_style = self.styles.get('subtitle', 'Normal')
        
        contact = cl_doc.add_paragraph(data.title)
        contact.style = subtitle_style
        if contact.runs: contact.runs[0].bold = True
        else: contact.add_run(data.title).bold = True
        
        details = cl_doc.add_paragraph(data.contact_info)
        details.style = subtitle_style

        cl_doc.add_paragraph()  # Spacer

        # --- DATE ---
        from datetime import datetime
        cl_doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        cl_doc.add_paragraph()

        # --- BODY ---
        # The prompt returns "Dear Hiring Manager,..." so we just dump it
        # We handle newlines by splitting
        for paragraph in letter_body.split('\n'):
            if paragraph.strip():
                p = cl_doc.add_paragraph(paragraph.strip())
                # Just use Normal style for letter body, maybe Justified if we want?
                # Sticking to Normal (Left aligned usually)
        
        cl_doc.add_paragraph()
        cl_doc.add_paragraph("Sincerely,")
        cl_doc.add_paragraph(data.name)

        cl_doc.save(output_filename)
        logger.info(f"Cover Letter generated successfully: {output_filename}")
